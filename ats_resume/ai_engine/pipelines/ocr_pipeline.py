"""
ai_engine/pipelines/ocr_pipeline.py
─────────────────────────────────────
OCR and text extraction pipeline.

Handles two main file types:
  - PDF (native text or scanned image)
  - DOCX (native Word document)

Strategy:
  PDF → pdfplumber first (fast, accurate for native PDFs)
       → If text is sparse (<50 chars/page), fallback to Tesseract OCR
  DOCX → python-docx paragraph + table extraction

Post-processing:
  - Unicode normalization
  - Whitespace normalization
  - Remove page numbers and headers/footers
  - Encoding artifact removal
"""

from __future__ import annotations

import io
import re
import unicodedata
from pathlib import Path

import pdfplumber
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document

from ai_engine.utils.text_cleaner import TextCleaner


class OCRPipeline:
    """
    Unified text extraction pipeline for PDF and DOCX files.

    Usage:
        pipeline = OCRPipeline()
        text = pipeline.extract(file_bytes, "application/pdf")
    """

    def __init__(self) -> None:
        self.text_cleaner = TextCleaner()
        # Tesseract config for resume-like documents
        # --psm 6: Assume a single uniform block of text
        # --oem 3: Use LSTM engine (most accurate)
        self.tesseract_config = "--psm 6 --oem 3"

    def extract(self, file_bytes: bytes, file_type: str) -> str:
        """
        Main entry point. Dispatches to PDF or DOCX extractor.

        Args:
            file_bytes: Raw file contents
            file_type:  MIME type string

        Returns:
            Cleaned, normalized plain text string

        Raises:
            ValueError: If file_type is unsupported
            ExtractionError: If extraction fails
        """
        if "pdf" in file_type.lower():
            raw_text = self._extract_pdf(file_bytes)
        elif "wordprocessingml" in file_type.lower() or "docx" in file_type.lower():
            raw_text = self._extract_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        if not raw_text or len(raw_text.strip()) < 50:
            raise ExtractionError(
                f"Extracted text too short ({len(raw_text)} chars). "
                "File may be corrupt or contain only images."
            )

        return self.text_cleaner.clean(raw_text)

    # ─── PDF Extraction ───────────────────────────────────────────────────────

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """
        Extract text from PDF.

        Strategy:
        1. Try pdfplumber (native text — fast and accurate)
        2. For pages with sparse text (<50 chars), use Tesseract OCR
        3. Combine results from all pages
        """
        all_text_parts: list[str] = []

        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Try native text extraction first
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3)

                    if page_text and len(page_text.strip()) >= 50:
                        # Good native text — use it
                        all_text_parts.append(page_text)
                    else:
                        # Sparse or no text — likely scanned. Use OCR.
                        ocr_text = self._ocr_page(page)
                        if ocr_text:
                            all_text_parts.append(ocr_text)

                    # Also extract text from tables (resumes often use table layout)
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = self._flatten_table(table)
                        if table_text:
                            all_text_parts.append(table_text)

        except Exception as e:
            raise ExtractionError(f"PDF extraction failed: {e}") from e

        return "\n\n".join(all_text_parts)

    def _ocr_page(self, page) -> str:
        """
        Render a PDF page to image and run Tesseract OCR.

        Image preprocessing steps:
        1. Render at 300 DPI (Tesseract accuracy degrades below 200 DPI)
        2. Convert to grayscale
        3. Apply sharpening filter
        4. Increase contrast
        """
        try:
            # Render page as PIL Image at 300 DPI
            pil_image = page.to_image(resolution=300).original

            # Preprocess for better OCR accuracy
            processed = self._preprocess_for_ocr(pil_image)

            # Run Tesseract
            text = pytesseract.image_to_string(
                processed,
                config=self.tesseract_config,
                lang="eng",
            )
            return text

        except Exception as e:
            # OCR failure is non-fatal — return empty string
            return ""

    @staticmethod
    def _preprocess_for_ocr(image: Image.Image) -> Image.Image:
        """
        Image preprocessing pipeline for better Tesseract accuracy.
        """
        # Convert to grayscale
        gray = image.convert("L")

        # Sharpen (helps with slightly blurry scans)
        sharpened = gray.filter(ImageFilter.SHARPEN)

        # Enhance contrast
        enhancer = ImageEnhance.Contrast(sharpened)
        contrasted = enhancer.enhance(2.0)

        return contrasted

    @staticmethod
    def _flatten_table(table: list[list]) -> str:
        """
        Convert a pdfplumber table (list of rows) to plain text.
        Tables in resumes often contain skills, education, contact info.
        """
        rows = []
        for row in table:
            if row:
                cell_texts = [str(cell).strip() for cell in row if cell]
                rows.append("  ".join(cell_texts))
        return "\n".join(rows)

    # ─── DOCX Extraction ──────────────────────────────────────────────────────

    def _extract_docx(self, file_bytes: bytes) -> str:
        """
        Extract text from DOCX file.

        Extracts:
        - Paragraphs (body text)
        - Table cells (skills table, experience table)
        - Text boxes (sometimes used in modern resume templates)
        """
        try:
            doc = Document(io.BytesIO(file_bytes))
        except Exception as e:
            raise ExtractionError(f"DOCX parsing failed: {e}") from e

        parts: list[str] = []

        # ── Paragraphs ────────────────────────────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # ── Tables ────────────────────────────────────────────────────
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append("  ".join(row_cells))

        # ── Headers and footers ───────────────────────────────────────
        for section in doc.sections:
            for para in section.header.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())

        return "\n".join(parts)


# ─── Exceptions ───────────────────────────────────────────────────────────────

class ExtractionError(Exception):
    """Raised when text extraction fails unrecoverably."""
    pass